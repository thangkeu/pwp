"""DocxAdapter — trích xuất từ file .docx dùng python-docx."""

from __future__ import annotations

import base64
import io
import zipfile

import docx
from docx.opc.exceptions import PackageNotFoundError

from app.adapters.base import ParserAdapter, ParsingError
from app.domain.parsed_document import (
    ParsedDocument,
    ParsedDocumentMetadata,
    ParsedImage,
    ParsedTable,
)


class DocxAdapter(ParserAdapter):
    @property
    def supported_extensions(self) -> list[str]:
        return ["docx"]

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        try:
            document = docx.Document(io.BytesIO(content))
        except (PackageNotFoundError, KeyError, ValueError, zipfile.BadZipFile) as exc:
            raise ParsingError(f"File DOCX không hợp lệ hoặc hỏng: {filename} ({exc})") from exc

        warnings: list[str] = []

        paragraphs_text = [p.text for p in document.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs_text)

        tables: list[ParsedTable] = []
        for table in document.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                tables.append(ParsedTable(rows=rows))

        headers: list[str] = []
        footers: list[str] = []
        for section in document.sections:
            header_text = "\n".join(p.text for p in section.header.paragraphs if p.text.strip())
            if header_text:
                headers.append(header_text)
            footer_text = "\n".join(p.text for p in section.footer.paragraphs if p.text.strip())
            if footer_text:
                footers.append(footer_text)

        comments: list[str] = []
        try:
            # python-docx không expose comments qua API chính thức trước 1.2 một cách đầy đủ;
            # đọc trực tiếp phần XML comments.xml nếu tồn tại trong package (best-effort).
            comments_part = document.part.package.part_related_by(
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
            )
            from lxml import etree

            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            tree = etree.fromstring(comments_part.blob)
            for c in tree.findall("w:comment", ns):
                comment_text = "".join(c.itertext()).strip()
                if comment_text:
                    comments.append(comment_text)
        except KeyError:
            pass  # Không có comment trong tài liệu — bình thường, không phải lỗi.
        except Exception as exc:  # noqa: BLE001 - best-effort, không chặn parse chính
            warnings.append(f"Không đọc được comments.xml: {exc}")

        images: list[ParsedImage] = []
        try:
            for idx, rel in enumerate(document.part.rels.values()):
                if "image" in rel.reltype:
                    image_bytes = rel.target_part.blob
                    images.append(
                        ParsedImage(
                            index=idx,
                            content_base64=base64.b64encode(image_bytes).decode("ascii"),
                        )
                    )
        except Exception as exc:  # noqa: BLE001
            warnings.append(f"Không trích xuất được toàn bộ hình ảnh: {exc}")

        word_count = len(text.split())

        return ParsedDocument(
            text=text,
            tables=tables,
            images=images,
            headers=headers,
            footers=footers,
            comments=comments,
            metadata=ParsedDocumentMetadata(
                parser_name="DocxAdapter",
                source_filename=filename,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                word_count=word_count,
                warnings=warnings,
            ),
        )
